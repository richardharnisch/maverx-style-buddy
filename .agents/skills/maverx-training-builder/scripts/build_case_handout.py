#!/usr/bin/env python3
"""Generate case_handout.docx for one session."""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from openrouter_call import call_openrouter

ROOT = Path(__file__).resolve().parent.parent


def write_case(intake: dict, track: dict, session: dict, prior_artefact: str) -> dict:
    system = (ROOT / "prompts" / "system_case_writer.md").read_text()
    user = (
        f"INTAKE:\n{json.dumps(intake, indent=2)}\n\n"
        f"TRACK:\n{json.dumps(track, indent=2)}\n\n"
        f"SESSION:\n{json.dumps(session, indent=2)}\n\n"
        f"PRIOR_ARTEFACT: {prior_artefact}\n"
    )
    return call_openrouter(system=system, user=user, max_tokens=3500, temperature=0.4)


def render(spec: dict, out_path: Path) -> None:
    doc = Document()
    title = doc.add_heading(spec.get("title", "Case handout"), level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x0D, 0x00, 0x6A)

    doc.add_heading("The case so far", level=1)
    doc.add_paragraph(spec.get("case_so_far", ""))

    doc.add_heading("Today's situation", level=1)
    doc.add_paragraph(spec.get("todays_situation", ""))

    for snip in spec.get("data_snippets", []) or []:
        doc.add_heading(snip.get("caption", "Data"), level=2)
        headers = snip.get("headers", [])
        rows = snip.get("rows", [])
        if headers:
            t = doc.add_table(rows=1 + len(rows), cols=len(headers))
            t.style = "Light List Accent 1"
            for ci, h in enumerate(headers):
                t.cell(0, ci).text = str(h)
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row[: len(headers)]):
                    t.cell(ri + 1, ci).text = str(val)

    doc.add_heading("Model reference", level=1)
    for m in spec.get("model_reference", []) or []:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{m.get('name', '')} — ").bold = True
        p.add_run(f"{m.get('purpose', '')} ({m.get('when_to_use', '')})")

    doc.add_heading("Working space", level=1)
    for w in spec.get("working_space", []) or []:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{w.get('prompt', '')}").bold = True
        if w.get("hint"):
            doc.add_paragraph(f"Hint: {w['hint']}")
        doc.add_paragraph("Your answer: _______________________________________________")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--intake", required=True)
    ap.add_argument("--track", required=True)
    ap.add_argument("--session", type=int, required=True)
    ap.add_argument("--out", required=True)
    args = ap.parse_args()
    intake = json.loads(Path(args.intake).read_text())
    track = json.loads(Path(args.track).read_text())
    session = next(s for s in track["sessions"] if s["n"] == args.session)
    prior = "none"
    if args.session > 1:
        prior = next(s for s in track["sessions"] if s["n"] == args.session - 1)["post_bite_artefact"]
    spec = write_case(intake, track, session, prior)
    render(spec, Path(args.out))
    print(f"wrote {args.out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
