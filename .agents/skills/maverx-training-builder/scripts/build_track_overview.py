#!/usr/bin/env python3
"""Generate track_overview.docx — the red-thread document for a Tier 3 track."""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

MAVERX_PURPLE = RGBColor(0x0D, 0x00, 0x6A)


def render(intake: dict, track: dict, out_path: Path) -> None:
    doc = Document()
    title = doc.add_heading(track.get("track_title", "Certification Track"), level=0)
    for r in title.runs:
        r.font.color.rgb = MAVERX_PURPLE

    doc.add_paragraph(
        f"Backbone: {track.get('backbone')}  |  "
        f"Sessions: {len(track.get('sessions', []))}  |  "
        f"Language: {track.get('language', intake.get('language', 'EN'))}"
    )

    doc.add_heading("The running case", level=1)
    case = track.get("case", {})
    company = case.get("company", {})
    doc.add_paragraph(
        f"Company: {company.get('name')} ({company.get('sector')}, "
        f"{company.get('size')}, {company.get('geography')})."
    )
    doc.add_paragraph(case.get("problem_statement", ""))

    doc.add_heading("Protagonists", level=2)
    for p in case.get("protagonists", []) or []:
        doc.add_paragraph(f"{p.get('name')} — {p.get('role')}", style="List Bullet")

    doc.add_heading("Datasets used in exercises", level=2)
    for tbl in case.get("dataset_spec", []) or []:
        doc.add_paragraph(tbl.get("table", ""), style="List Bullet")
        for col in tbl.get("columns", []) or []:
            doc.add_paragraph(
                f"  {col.get('name')} ({col.get('type')}) {col.get('range', '')}",
                style="List Bullet 2",
            )

    doc.add_heading("The red thread (session by session)", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Phase"
    hdr[2].text = "Title"
    hdr[3].text = "Objective"
    hdr[4].text = "Post-bite artefact"
    for s in track.get("sessions", []):
        row = t.add_row().cells
        row[0].text = str(s.get("n"))
        row[1].text = str(s.get("backbone_phase", ""))
        row[2].text = str(s.get("title", ""))
        row[3].text = str(s.get("objective", ""))
        row[4].text = str(s.get("post_bite_artefact", ""))

    doc.add_heading("Handshake between sessions", level=1)
    sessions = track.get("sessions", [])
    for i in range(1, len(sessions)):
        prev = sessions[i - 1]
        curr = sessions[i]
        ok = prev.get("post_bite_artefact") == curr.get("next_session_pre_bite_expects_from_prior")
        mark = "OK" if ok else "MISMATCH"
        doc.add_paragraph(
            f"[{mark}] Session {prev['n']} produces {prev.get('post_bite_artefact')!r}; "
            f"Session {curr['n']} pre-bite expects {curr.get('next_session_pre_bite_expects_from_prior')!r}.",
            style="List Bullet",
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--intake", required=True)
    ap.add_argument("--track", required=True)
    ap.add_argument("--out", required=True)
    args = ap.parse_args()
    intake = json.loads(Path(args.intake).read_text())
    track = json.loads(Path(args.track).read_text())
    render(intake, track, Path(args.out))
    print(f"wrote {args.out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
