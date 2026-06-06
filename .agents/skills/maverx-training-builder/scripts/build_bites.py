#!/usr/bin/env python3
"""Generate pre_bite.docx and post_bite.docx for a session."""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from openrouter_call import call_openrouter, DEFAULT_WRITER_MODEL

ROOT = Path(__file__).resolve().parent.parent


def write_bite(kind: str, intake: dict, track: dict, session: dict) -> dict:
    system = (ROOT / "prompts" / "system_bite_writer.md").read_text()
    user = (
        f"KIND: {kind}\n"
        f"INTAKE:\n{json.dumps(intake, indent=2)}\n\n"
        f"TRACK:\n{json.dumps(track, indent=2)}\n\n"
        f"SESSION:\n{json.dumps(session, indent=2)}\n"
    )
    return call_openrouter(system=system, user=user, max_tokens=2500, temperature=0.4)


def render(doc_spec: dict, out_path: Path) -> None:
    doc = Document()
    # Title
    title = doc.add_heading(doc_spec.get("title", ""), level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x00, 0x6A)
    for section in doc_spec.get("sections", []):
        doc.add_heading(section["heading"], level=1)
        p = doc.add_paragraph(section["body"])
        for run in p.runs:
            run.font.size = Pt(11)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--intake", required=True)
    ap.add_argument("--track", required=True)
    ap.add_argument("--session", type=int, required=True)
    ap.add_argument("--out-dir", required=True)
    args = ap.parse_args()

    intake = json.loads(Path(args.intake).read_text())
    track = json.loads(Path(args.track).read_text())
    session = next(s for s in track["sessions"] if s["n"] == args.session)
    out_dir = Path(args.out_dir)

    for kind, fname in [("pre_bite", "pre_bite.docx"), ("post_bite", "post_bite.docx")]:
        spec = write_bite(kind, intake, track, session)
        render(spec, out_dir / fname)
        print(f"wrote {out_dir / fname}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
